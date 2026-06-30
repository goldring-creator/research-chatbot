# -*- coding: utf-8 -*-
"""
app.py — 사회과학 연구설계 챗봇 (데스크탑 앱)
NVIDIA 무료 API + 모래시계 연구 로직. 터미널 느낌의 플로팅 채팅창.
각 사용자가 본인의 무료 NVIDIA 키를 입력해 사용한다(키는 본인 컴퓨터에만 저장).
"""

import os
import re
import json
import queue
import threading
import platform
import webbrowser
import subprocess
from datetime import datetime

import tkinter as tk
from tkinter import filedialog, messagebox

import core
import prompts

# ── 데이터 저장 위치 ──
DATA_DIR = os.path.expanduser("~/Documents/ResearchChatbot")
OUT_DIR = os.path.join(DATA_DIR, "outputs")
CONV_DIR = os.path.join(DATA_DIR, "conversations")
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(CONV_DIR, exist_ok=True)

# ── 어두운 테마 + 밝은 테두리 ──
C_BG = "#0d1117"        # 전체 배경(어두움)
C_BG2 = "#161b22"       # 패널(툴바·입력창)
C_BORDER = "#30363d"    # 은은한 구분선
C_LINE = "#cdd5df"      # 카드 테두리(밝은/흰색 느낌)
C_TEXT = "#e6edf3"      # 본문 글자(밝음)
C_DIM = "#8b949e"       # 흐린 글자
C_GREEN = "#3fb950"     # 강조(버튼 배경)
C_GREEN_D = "#2ea043"   # 버튼 hover
C_GREEN_TX = "#3fb950"  # 초록 글자(어두운 배경 위)
C_BLUE = "#58a6ff"
C_AMBER = "#e3b341"
C_RED = "#f85149"
C_HR = "#30363d"

# 폰트: 터미널 느낌의 고정폭
if platform.system() == "Darwin":
    MONO = "Menlo"
elif platform.system() == "Windows":
    MONO = "Consolas"
else:
    MONO = "DejaVu Sans Mono"

NVIDIA_URL = "https://build.nvidia.com"

# ── 모드별 활용 안내 (모드 전환 시 표시) ──
MODE_DESC = {
    "design": [
        "연구 아이디어 → 서론(배경·갭·RQ)·이론적 틀·표집·변수·측정·분석 설계",
        "예상 타당도 위협과 통제 방안 제안",
        "예) \"교사 AI 활용과 직무만족 관계를 연구하고 싶어\"",
    ],
    "review": [
        "초안을 입력하거나 📎로 파일 첨부 → 모래시계 구조로 검토",
        "결과·해석 분리, 논의 6단계, 재현가능성, 인용 점검",
        "예) 📎 내 서론 초안.hwpx 첨부 → 피드백",
    ],
    "chat": [
        "연구방법론·통계·연구설계에 대한 자유 질문",
        "개념 설명, 연구 간 비교, 분석방법 추천 등",
        "예) \"위계적 회귀와 SEM 중 뭐가 적합해?\"",
    ],
}

# ── 버전·업데이트 확인 ──
APP_VERSION = "0.1.7"
REPO = "goldring-creator/research-chatbot"
RELEASES_URL = f"https://github.com/{REPO}/releases/latest"

# ── 키 발급 안내 (1) 2) 3) 순서) ──
GUIDE_STEPS = [
    "1)  아래 [NVIDIA 사이트 열기]를 눌러 로그인 (구글·이메일)",
    "2)  화면 위에 Verify가 보이면 눌러 계정 인증",
    "3)  아무 모델이나 열기 (예: 검색창에 deepseek)",
    "4)  Build 탭의 [Generate API Key] 클릭",
    "5)  만들어진  nvapi-...  키를 [Copy]",
    "6)  아래 칸에 붙여넣고 [저장하고 시작]",
]


class _Tooltip:
    """아이콘에 커서를 올리면 설명을 띄운다 (Mac·Windows 공통)."""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip = None
        widget.bind("<Enter>", self.show, add="+")
        widget.bind("<Leave>", self.hide, add="+")

    def show(self, e=None):
        if self.tip:
            return
        x = self.widget.winfo_rootx() + 6
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 4
        root = self.widget.winfo_toplevel()
        self.tip = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        try:
            tw.wm_attributes("-topmost", True)
        except Exception:
            pass
        tk.Label(tw, text=self.text, bg="#1f2630", fg="#e6edf3",
                 font=(MONO, 9), padx=8, pady=4, relief="solid", bd=1).pack()
        tw.wm_geometry(f"+{x}+{y}")
        tw.update_idletasks()

        def raise_tip():
            if not tw.winfo_exists():
                return
            try:
                tw.lift(root)                 # 메인(항상 위) 창 바로 위로 명시적 올림
                tw.wm_attributes("-topmost", True)
            except Exception:
                pass
        raise_tip()
        tw.after(20, raise_tip)               # 창 매니저가 재배치한 뒤 한 번 더

    def hide(self, e=None):
        if self.tip:
            self.tip.destroy()
            self.tip = None


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("연구설계 챗봇")
        self.configure(bg=C_BG)
        self.geometry("640x720")
        self.minsize(600, 560)
        self.attributes("-topmost", True)
        self.pinned = True

        self.mode = "design"
        self.model_key = "deepseek"
        self.history = []
        self.client = None
        self.q = queue.Queue()
        self.streaming = False
        self._cancel = False          # 사용자가 답변/내보내기 중지를 눌렀는지
        self._export_kind = None      # 내보내기 진행 중이면 "html"/"docx"

        key = core.load_key()
        if core.valid_key_format(key):
            self.client = core.make_client(key)
            self.build_main()
        else:
            self.build_onboarding()

        self.setup_hotkey()

    # ── 복사·붙여넣기 단축키 활성화 (macOS Tk 보완) ──
    def _enable_clipboard(self, w):
        for mod in ("Command", "Control"):
            w.bind(f"<{mod}-v>", lambda e: (e.widget.event_generate("<<Paste>>"), "break")[1])
            w.bind(f"<{mod}-c>", lambda e: (e.widget.event_generate("<<Copy>>"), "break")[1])
            w.bind(f"<{mod}-x>", lambda e: (e.widget.event_generate("<<Cut>>"), "break")[1])
            w.bind(f"<{mod}-a>", self._select_all)

    def _select_all(self, e):
        w = e.widget
        try:
            w.select_range(0, "end")        # Entry
        except Exception:
            w.tag_add("sel", "1.0", "end-1c")  # Text
        return "break"

    # ── 라벨 기반 버튼 (macOS에서도 색이 제대로 나옴) ──
    def _btn(self, parent, text, cmd, fg=C_TEXT, bg=C_BG2, hover="#21262d",
             font=None, pady=6, padx=10):
        lbl = tk.Label(parent, text=text, fg=fg, bg=bg, cursor="hand2",
                       font=font or (MONO, 10), padx=padx, pady=pady)
        lbl.bind("<Button-1>", lambda e: cmd())
        lbl.bind("<Enter>", lambda e: lbl.configure(bg=hover))
        lbl.bind("<Leave>", lambda e: lbl.configure(bg=bg))
        return lbl

    # ════════ 온보딩 (키 입력) ════════
    def build_onboarding(self):
        for w in self.winfo_children():
            w.destroy()

        outer = tk.Frame(self, bg=C_BG, padx=26, pady=20)
        outer.pack(fill="both", expand=True)

        # 맨 아래: 저장 버튼 (먼저 bottom에 배치)
        bottom = tk.Frame(outer, bg=C_BG)
        bottom.pack(side="bottom", fill="x")
        save_btn = tk.Label(bottom, text="저장하고 시작  ▸", bg=C_GREEN, fg="#0d1117",
                            font=(MONO, 14, "bold"), cursor="hand2", pady=12)
        save_btn.pack(fill="x", pady=(10, 0))
        save_btn.bind("<Button-1>", lambda e: self.save_and_start())
        save_btn.bind("<Enter>", lambda e: save_btn.configure(bg=C_GREEN_D))
        save_btn.bind("<Leave>", lambda e: save_btn.configure(bg=C_GREEN))

        # 위: 제목 + 안내 + 키 입력
        top = tk.Frame(outer, bg=C_BG)
        top.pack(side="top", fill="both", expand=True)

        # 제목 (가운데, 큰 폰트)
        tk.Label(top, text="🧭  사회과학 연구설계 챗봇", bg=C_BG, fg=C_GREEN_TX,
                 font=(MONO, 19, "bold")).pack(pady=(6, 2))
        tk.Label(top, text="NVIDIA 무료 AI · 본인 키로 작동 · 완전 무료",
                 bg=C_BG, fg=C_DIM, font=(MONO, 10)).pack(pady=(0, 16))

        # 키 발급 방법
        tk.Label(top, text="📋  무료 API 키 발급 방법 (약 3분)", bg=C_BG, fg=C_BLUE,
                 font=(MONO, 12, "bold")).pack(anchor="w")
        steps = tk.Frame(top, bg=C_BG2)
        steps.pack(fill="x", pady=(6, 10))
        for s in GUIDE_STEPS:
            tk.Label(steps, text=s, bg=C_BG2, fg=C_TEXT, font=(MONO, 11),
                     justify="left", anchor="w", wraplength=400,
                     padx=12, pady=3).pack(fill="x")

        # NVIDIA 사이트 열기 (하이퍼링크 버튼)
        link = tk.Label(top, text="🔗  NVIDIA 사이트 열기  (build.nvidia.com)",
                        bg=C_BG, fg=C_BLUE, font=(MONO, 11, "underline"),
                        cursor="hand2")
        link.pack(anchor="w", pady=(0, 16))
        link.bind("<Button-1>", lambda e: webbrowser.open(NVIDIA_URL))

        # 키 입력칸
        tk.Label(top, text="🔑  발급받은 API 키 붙여넣기", bg=C_BG, fg=C_TEXT,
                 font=(MONO, 11, "bold")).pack(anchor="w")
        self.key_entry = tk.Entry(top, bg="#ffffff", fg="#1A1A1A",
                                  insertbackground="#1A1A1A", font=(MONO, 12),
                                  relief="flat", highlightthickness=1,
                                  highlightbackground=C_BORDER, highlightcolor=C_BLUE)
        self.key_entry.pack(fill="x", ipady=8, pady=(6, 4))
        self.key_entry.focus_set()
        self._enable_clipboard(self.key_entry)
        tk.Label(top, text="🔒 이 키는 당신 컴퓨터에만 저장되며 외부로 전송되지 않습니다.",
                 bg=C_BG, fg=C_DIM, font=(MONO, 9), wraplength=400,
                 justify="left").pack(anchor="w")

    def save_and_start(self):
        key = self.key_entry.get().strip()
        if not core.valid_key_format(key):
            messagebox.showerror("키 형식 오류",
                                 "키는 nvapi- 로 시작하는 영문/숫자여야 합니다.\n"
                                 "한글이나 빈칸이 섞이지 않았는지 확인하세요.")
            return
        core.save_key(key)
        self.client = core.make_client(key)
        self.build_main()

    # ════════ 메인 채팅 화면 ════════
    def build_main(self):
        for w in self.winfo_children():
            w.destroy()

        bar = tk.Frame(self, bg=C_BG2)
        bar.pack(fill="x")
        self._toolbar = bar
        left = tk.Frame(bar, bg=C_BG2)
        left.pack(side="left", padx=8, pady=6)
        right = tk.Frame(bar, bg=C_BG2)
        right.pack(side="right", padx=8, pady=6)

        # 좌측: 모드 / 모델
        self.mode_btn = self._btn(left, self._mode_label(), self.open_mode_menu, fg=C_GREEN_TX)
        self.mode_btn.pack(side="left", padx=4)
        _Tooltip(self.mode_btn, "모드 선택 — 클릭하면 설계 / 검토 / 자유문답 목록이 열립니다")

        # 우측: 아이콘 (균등 간격)
        for txt, cmd, tip in [
            ("📎", self.attach_file, "파일 첨부(여러 개 가능) — 초안을 검토 모드로 분석"),
            ("📤", self.open_export_menu, "지금까지 내용을 HTML·Word 문서로 내보내기"),
            ("💾", self.save_conv, "현재 대화 저장"),
            ("📂", self.load_conv, "저장한 대화 불러오기"),
            ("📌", self.toggle_pin, "항상 위에 고정 켜기/끄기"),
            ("🔑", self.change_key, "API 키 변경"),
        ]:
            b = self._btn(right, txt, cmd, padx=9)
            b.pack(side="left", padx=5)
            _Tooltip(b, tip)
        tk.Frame(self, bg=C_BORDER, height=1).pack(side="top", fill="x")  # 툴바 구분선

        # 각주: 읽기 지원 형식 (맨 아래)
        foot = tk.Label(self, text="📎 " + core.SUPPORTED_NOTE, bg=C_BG, fg=C_DIM,
                        font=(MONO, 9))
        foot.pack(side="bottom", fill="x", pady=(0, 6))

        # 입력 영역 — 맨 아래 고정 + 안쪽 카드(밝은 테두리)
        ibar_outer = tk.Frame(self, bg=C_BG)
        ibar_outer.pack(side="bottom", fill="x", padx=10, pady=(4, 0))
        ibar = tk.Frame(ibar_outer, bg=C_BG2, highlightthickness=1,
                        highlightbackground=C_LINE, highlightcolor=C_LINE, bd=0)
        ibar.pack(fill="x")
        # 보내기 버튼을 먼저 오른쪽에 예약 → 좁아져도 안 잘림
        self.send_btn = tk.Label(ibar, text="보내기", bg=C_GREEN, fg="#0d1117",
                                 font=(MONO, 11, "bold"), cursor="hand2", padx=14)
        self.send_btn.pack(side="right", fill="y", padx=(4, 6), pady=6)
        self.send_btn.bind("<Button-1>", lambda e: self._on_send_click())
        self.bind("<Escape>", lambda e: self.cancel_stream())   # Esc로도 중지
        tk.Label(ibar, text=" 입력 ▸", bg=C_BG2, fg=C_GREEN_TX,
                 font=(MONO, 11, "bold")).pack(side="left", anchor="n", pady=12)
        self.inp = tk.Text(ibar, bg=C_BG2, fg=C_TEXT, font=(MONO, 11), height=3,
                          wrap="word", relief="flat", bd=0, padx=6, pady=8,
                          insertbackground=C_TEXT)
        self.inp.pack(side="left", fill="both", expand=True)
        self.inp.bind("<Return>", self.on_return)
        self.inp.bind("<FocusIn>", self._clear_placeholder)
        self.inp.bind("<FocusOut>", self._restore_placeholder)
        self._enable_clipboard(self.inp)
        self._ph_active = False
        self._set_placeholder()

        # 대화 영역 — 창 안쪽으로 들인 카드(밝은 테두리)
        wrap = tk.Frame(self, bg=C_BG)
        wrap.pack(side="top", fill="both", expand=True, padx=10, pady=(8, 4))
        card = tk.Frame(wrap, bg=C_BG, highlightthickness=1,
                        highlightbackground=C_LINE, highlightcolor=C_LINE, bd=0)
        card.pack(fill="both", expand=True)
        self.chat = tk.Text(card, bg=C_BG, fg=C_TEXT, font=(MONO, 11), wrap="word",
                            relief="flat", bd=0, padx=12, pady=10, insertbackground=C_TEXT,
                            state="disabled", spacing1=2, spacing3=4)
        self._sb = tk.Scrollbar(card, command=self.chat.yview)
        self.chat.configure(yscrollcommand=self._on_chat_scroll)
        self._sb.pack(side="right", fill="y")
        self.chat.pack(side="left", fill="both", expand=True)
        self._enable_copy(self.chat)   # 답변 드래그 선택·복사·전체 선택 허용
        self.chat.tag_config("user", foreground=C_BLUE, font=(MONO, 11, "bold"))
        self.chat.tag_config("bot", foreground=C_TEXT)
        self.chat.tag_config("sys", foreground=C_DIM, font=(MONO, 10))
        self.chat.tag_config("warn", foreground=C_AMBER, font=(MONO, 10))
        self.chat.tag_config("err", foreground=C_RED)
        self.chat.tag_config("label", foreground=C_GREEN_TX, font=(MONO, 10, "bold"))
        self.chat.tag_config("b", font=(MONO, 11, "bold"))
        self.chat.tag_config("h", foreground=C_GREEN_TX, font=(MONO, 12, "bold"))
        self.chat.tag_config("hr", foreground=C_HR)
        self.chat.tag_config("code", foreground="#c9d1d9", font=(MONO, 10))

        self._sys(f"준비 완료. 모드: {prompts.MODE_LABELS[self.mode]} · "
                  f"모델: {core.MODEL_LABELS[self.model_key]}")
        self._sys("연구 아이디어를 입력하거나 📎로 초안을 첨부하세요. "
                  "(Enter 전송 / Shift+Enter 줄바꿈)")
        self._mode_help()      # 현재 모드로 가능한 작업 안내
        self._check_update()   # 백그라운드로 새 버전 확인

    # ════════ 업데이트 알림 ════════
    def _check_update(self):
        def work():
            try:
                import urllib.request
                req = urllib.request.Request(
                    f"https://api.github.com/repos/{REPO}/releases/latest",
                    headers={"User-Agent": "research-chatbot"})
                with urllib.request.urlopen(req, timeout=8) as r:
                    data = json.load(r)
                latest = (data.get("tag_name") or "").lstrip("v")
                if latest and self._is_newer(latest, APP_VERSION):
                    self.after(0, lambda: self._show_update(latest))
            except Exception:
                pass  # 네트워크 실패 시 조용히 무시
        threading.Thread(target=work, daemon=True).start()

    def _is_newer(self, a, b):
        pa = [int(x) for x in re.findall(r"\d+", a)]
        pb = [int(x) for x in re.findall(r"\d+", b)]
        return pa > pb

    def _show_update(self, latest):
        if getattr(self, "_update_shown", False):
            return
        self._update_shown = True
        banner = tk.Label(self, text=f"🆕 새 버전 v{latest} 이(가) 있습니다 — 클릭해서 다운로드",
                          bg="#1f6feb", fg="#ffffff", font=(MONO, 10, "bold"),
                          cursor="hand2", pady=6)
        banner.pack(side="top", fill="x", before=self._toolbar)
        banner.bind("<Button-1>", lambda e: webbrowser.open(RELEASES_URL))

    def _mode_label(self):
        return f" 모드:{prompts.MODE_LABELS[self.mode]} ▾ "

    # ════════ 토글/명령 ════════
    MODE_ORDER = ["design", "review", "chat"]

    def open_mode_menu(self):
        """모드 버튼을 누르면 선택 목록을 띄운다(클릭해서 직접 고르기)."""
        menu = tk.Menu(self, tearoff=0, bg=C_BG2, fg=C_TEXT,
                       activebackground=C_GREEN, activeforeground="#0d1117",
                       font=(MONO, 10), bd=0)
        for key in self.MODE_ORDER:
            mark = "● " if key == self.mode else "○ "
            menu.add_command(label=mark + prompts.MODE_LABELS[key],
                             command=lambda k=key: self.set_mode(k))
        x = self.mode_btn.winfo_rootx()
        y = self.mode_btn.winfo_rooty() + self.mode_btn.winfo_height()
        try:
            menu.tk_popup(x, y)
        finally:
            menu.grab_release()

    def set_mode(self, key):
        if key == self.mode:
            return
        self.mode = key
        self.mode_btn.configure(text=self._mode_label())
        self._sys(f"→ {prompts.MODE_LABELS[self.mode]} 모드로 변경")
        self._mode_help()

    # ── 보내기/중지 버튼 상태 ──
    def _set_stop_button(self):
        self.send_btn.configure(text="■ 중지", bg=C_RED, fg="#0d1117")

    def _set_send_button(self):
        self.send_btn.configure(text="보내기", bg=C_GREEN, fg="#0d1117")

    def _on_send_click(self):
        if self.streaming:
            self.cancel_stream()
        else:
            self.send()

    def cancel_stream(self):
        """진행 중인 답변/내보내기를 중지한다(부분 결과는 유지)."""
        if not self.streaming:
            return
        self._cancel = True
        self.send_btn.configure(text="중지 중…")

    def _mode_help(self):
        for line in MODE_DESC.get(self.mode, []):
            self._append("   • " + line + "\n", "sys")

    def toggle_pin(self):
        self.pinned = not self.pinned
        self.attributes("-topmost", self.pinned)
        self._sys("항상 위: " + ("켜짐 📌" if self.pinned else "꺼짐"))

    def change_key(self):
        if messagebox.askyesno("키 변경", "API 키를 다시 입력하시겠습니까?"):
            self.build_onboarding()

    def show_help(self):
        win = tk.Toplevel(self, bg=C_BG)
        win.title("도움말 — 키 발급 방법")
        win.configure(padx=22, pady=18)
        win.attributes("-topmost", True)
        tk.Label(win, text="무료 NVIDIA API 키 발급 방법", bg=C_BG, fg=C_GREEN_TX,
                 font=(MONO, 13, "bold")).pack(anchor="w", pady=(0, 8))
        for s in GUIDE_STEPS:
            tk.Label(win, text=s, bg=C_BG, fg=C_TEXT, font=(MONO, 11),
                     justify="left", wraplength=440).pack(anchor="w", pady=1)
        link = tk.Label(win, text="🔗 build.nvidia.com 열기", bg=C_BG, fg=C_BLUE,
                        font=(MONO, 11, "underline"), cursor="hand2")
        link.pack(anchor="w", pady=(8, 0))
        link.bind("<Button-1>", lambda e: webbrowser.open(NVIDIA_URL))

    def attach_file(self):
        paths = filedialog.askopenfilenames(
            title="검토할 파일 선택 (여러 개 가능)",
            filetypes=[("지원 문서", "*.txt *.md *.pdf *.docx"), ("모든 파일", "*.*")])
        if not paths:
            return
        parts, names = [], []
        for p in paths:
            content, err = core.read_file(p)
            if err:
                self._warn(f"{os.path.basename(p)}: {err}")
                continue
            parts.append(f"[파일: {os.path.basename(p)}]\n{content}")
            names.append(os.path.basename(p))
        if not parts:
            return
        self.mode = "review"
        self.mode_btn.configure(text=self._mode_label())
        self._sys(f"📎 {len(names)}개 파일 읽음 → 검토 모드: {', '.join(names)}")
        self._dispatch("다음 원고를 검토해줘:\n\n" + "\n\n".join(parts),
                       display=f"[첨부] {', '.join(names)} 검토 요청")

    # ════════ 내보내기 (HTML / Word) ════════
    def open_export_menu(self):
        """📤 — 클릭 시 HTML / Word 중 형식을 고르는 메뉴를 띄운다."""
        if self.streaming:
            self._sys("답변 생성이 끝난 뒤 내보내기를 해주세요.")
            return
        if not any(m["role"] == "assistant" for m in self.history):
            self._sys("먼저 내용을 작성한 뒤 눌러주세요 (설계·검토·문답 등).")
            return
        menu = tk.Menu(self, tearoff=0, bg=C_BG2, fg=C_TEXT,
                       activebackground=C_GREEN, activeforeground="#0d1117",
                       font=(MONO, 10), bd=0)
        menu.add_command(label="HTML 문서 (.html)",
                         command=lambda: self.start_export("html"))
        menu.add_command(label="Word 문서 (.docx)",
                         command=lambda: self.start_export("docx"))
        try:
            menu.tk_popup(self.winfo_pointerx(), self.winfo_pointery())
        finally:
            menu.grab_release()

    def start_export(self, kind):
        """선택한 형식으로 '지금까지 내용'을 문서로 생성한다(대화창·기록은 건드리지 않음)."""
        if self.streaming:
            return
        instruction = (prompts.EXPORT_HTML_INSTRUCTION if kind == "html"
                       else prompts.EXPORT_DOCX_INSTRUCTION)
        msgs = self.history + [{"role": "user", "content": instruction}]   # history 미변경
        self._export_kind = kind
        self._exp_acc = []
        self._exp_err = None
        self._cancel = False
        self.streaming = True
        self._set_stop_button()
        label = "HTML" if kind == "html" else "Word"
        self._sys(f"📤 {label} 문서 생성 중…  (중지하려면 ■ 중지 또는 Esc)")
        self._exp_start = self.chat.index("end-1c")
        self._exp_load_i = 0
        self._retry_msg = None
        threading.Thread(target=self._export_worker, args=(msgs,), daemon=True).start()
        self.after(40, self._export_poll)
        self._animate_export()

    def _export_worker(self, msgs):
        for piece in core.stream_answer(self.client, self.model_key, "code", msgs,
                                        should_cancel=lambda: self._cancel):
            if isinstance(piece, tuple) and piece[0] == "__error__":
                self.q.put(("err", piece[1]))
            elif isinstance(piece, tuple) and piece[0] == "__retry__":
                self.q.put(("retry", piece[1:]))
            else:
                self.q.put(("chunk", piece))
        self.q.put(("done", None))

    def _export_poll(self):
        try:
            while True:
                kind, data = self.q.get_nowait()
                if kind == "chunk":
                    self._exp_acc.append(data)        # 화면엔 출력하지 않음
                elif kind == "retry":
                    wait, attempt, total = data
                    self._retry_msg = (
                        f"요청이 많아 약 {wait}초 후 자동으로 다시 시도합니다 (재시도 {attempt}/{total})"
                    )
                elif kind == "err":
                    self._exp_err = data
                elif kind == "done":
                    self._finish_export()
                    return
        except queue.Empty:
            pass
        if self.streaming and self._export_kind is not None:
            self.after(40, self._export_poll)

    def _animate_export(self):
        if not self.streaming or self._export_kind is None:
            return
        retry = getattr(self, "_retry_msg", None)
        base = ("   " + retry + " ") if retry else "   생성중 "
        frames = ["·", "· ·", "· · ·", "· · · ·"]
        follow = self._at_bottom()
        self.chat.configure(state="normal")
        self.chat.delete(self._exp_start, "end-1c")
        self.chat.insert("end", base + frames[self._exp_load_i % len(frames)], "sys")
        if follow:
            self.chat.see("end")
        self.chat.configure(state="disabled")
        self._exp_load_i += 1
        self.after(400, self._animate_export)

    def _finish_export(self):
        self.streaming = False
        self._set_send_button()
        # 스피너 줄 제거
        self.chat.configure(state="normal")
        self.chat.delete(self._exp_start, "end-1c")
        self.chat.configure(state="disabled")
        kind = self._export_kind
        self._export_kind = None
        if self._cancel:
            self._cancel = False
            self._sys("내보내기를 취소했습니다.")
            return
        if self._exp_err:
            self._warn(core.friendly_error(self._exp_err))
            self._exp_err = None
            return
        answer = "".join(self._exp_acc)
        if not answer.strip():
            self._warn("문서를 생성하지 못했습니다. 다시 시도해 주세요.")
            return
        if kind == "html":
            self._save_html(answer)
        else:
            self._save_docx(answer)
        self._notify("문서 생성 완료", "문서 생성이 완료되었습니다.")

    def _open_file(self, path):
        """저장한 파일을 기본 프로그램으로 연다(브라우저·Word 등)."""
        try:
            if platform.system() == "Darwin":
                subprocess.run(["open", path], check=False)
            elif platform.system() == "Windows":
                os.startfile(path)  # type: ignore[attr-defined]
            else:
                subprocess.run(["xdg-open", path], check=False)
        except Exception:
            webbrowser.open("file://" + path)   # 최후 수단

    def _save_html(self, answer):
        html = answer.strip()
        m = re.search(r"<!DOCTYPE html.*?</html\s*>", html, re.S | re.I)
        if m:
            html = m.group(0)
        else:
            html = re.sub(r"^```[a-zA-Z]*\n?", "", html)
            html = re.sub(r"\n?```$", "", html.strip())
        ts = datetime.now().strftime("%Y-%m-%d_%H%M")
        path = filedialog.asksaveasfilename(
            title="HTML 저장 위치·파일명 선택",
            initialdir=OUT_DIR, initialfile=f"{ts}_생성문서.html",
            defaultextension=".html",
            filetypes=[("HTML 문서", "*.html"), ("모든 파일", "*.*")])
        if not path:
            self._sys("내보내기를 취소했습니다.")
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(html)
            self._sys(f"📤 HTML 저장됨 — 엽니다:\n   {path}")
            self._open_file(path)
        except Exception as e:
            self._warn(f"HTML 저장 오류: {e}")

    def _save_docx(self, answer):
        """마크다운 형태의 답변을 .docx로 변환해 저장·연다(제목·굵게·목록·표 지원)."""
        try:
            import docx
        except Exception:
            self._warn("Word 내보내기에는 python-docx가 필요합니다.")
            return
        text = answer.strip()
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)   # 혹시 모를 코드펜스 제거
        text = re.sub(r"\n?```$", "", text.strip())
        doc = docx.Document()
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line.strip():
                i += 1
                continue
            # 표 블록 (연속된 | … | 줄)
            if line.lstrip().startswith("|") and line.rstrip().endswith("|"):
                rows = []
                while i < len(lines) and lines[i].lstrip().startswith("|"):
                    rows.append(lines[i].strip())
                    i += 1
                self._docx_table(doc, rows)
                continue
            m = re.match(r"^(#{1,6})\s*(.*)", line)
            if m:                                       # 제목
                doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
                i += 1
                continue
            mb = re.match(r"^\s*[-*]\s+(.*)", line)
            if mb:                                      # 글머리 목록
                self._docx_runs(doc.add_paragraph(style="List Bullet"), mb.group(1))
                i += 1
                continue
            self._docx_runs(doc.add_paragraph(), line)  # 일반 단락
            i += 1
        ts = datetime.now().strftime("%Y-%m-%d_%H%M")
        path = filedialog.asksaveasfilename(
            title="Word 저장 위치·파일명 선택",
            initialdir=OUT_DIR, initialfile=f"{ts}_생성문서.docx",
            defaultextension=".docx",
            filetypes=[("Word 문서", "*.docx"), ("모든 파일", "*.*")])
        if not path:
            self._sys("내보내기를 취소했습니다.")
            return
        try:
            doc.save(path)
            self._sys(f"📤 Word 저장됨 — 엽니다:\n   {path}")
            self._open_file(path)
        except Exception as e:
            self._warn(f"Word 저장 오류: {e}")

    def _docx_runs(self, paragraph, text):
        """단락 안의 **굵게**를 굵은 run으로 나눠 넣는다(남은 마크다운 기호는 제거)."""
        text = text.replace("`", "")
        for part in re.split(r"(\*\*.+?\*\*)", text):
            if len(part) >= 4 and part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            elif part:
                paragraph.add_run(part.replace("**", ""))

    def _docx_table(self, doc, rows):
        """마크다운 표 행들을 docx 표로 만든다(첫 행=헤더, 구분행 ---는 무시)."""
        def cells(r):
            return [c.strip() for c in r.strip().strip("|").split("|")]
        data = [cells(r) for r in rows
                if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]   # |---|---| 제거
        if not data:
            return
        ncol = max(len(r) for r in data)
        table = doc.add_table(rows=0, cols=ncol)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for ri, row in enumerate(data):
            cellrow = table.add_row().cells
            for ci in range(ncol):
                para = cellrow[ci].paragraphs[0]
                self._docx_runs(para, row[ci] if ci < len(row) else "")
                if ri == 0:                              # 헤더 굵게
                    for run in para.runs:
                        run.bold = True
        doc.add_paragraph()                              # 표 뒤 간격

    def save_conv(self):
        if not self.history:
            self._sys("저장할 대화가 없습니다."); return
        ts = datetime.now().strftime("%Y-%m-%d_%H%M%S")
        path = filedialog.asksaveasfilename(
            title="대화 저장 — 위치·파일명 선택",
            initialdir=CONV_DIR, initialfile=f"conv_{ts}.json",
            defaultextension=".json",
            filetypes=[("대화 파일", "*.json"), ("모든 파일", "*.*")])
        if not path:
            self._sys("저장을 취소했습니다."); return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.history, f, ensure_ascii=False, indent=2)
            self._sys(f"💾 저장됨: {path}")
        except Exception as e:
            self._warn(f"저장 오류: {e}")

    def load_conv(self):
        path = filedialog.askopenfilename(initialdir=CONV_DIR, title="대화 불러오기",
                                          filetypes=[("대화", "*.json")])
        if not path:
            return
        with open(path, "r", encoding="utf-8") as f:
            self.history = json.load(f)
        self.chat.configure(state="normal"); self.chat.delete("1.0", "end")
        self.chat.configure(state="disabled")
        for m in self.history:
            if m["role"] == "user":
                self._append("\n나 ▸ ", "user"); self._append(m["content"] + "\n", "bot")
            else:
                self._append("\nAI ▸\n", "label"); self._append(m["content"] + "\n", "bot")
        self._sys(f"📂 불러옴 (메시지 {len(self.history)}개)")

    # ════════ 입력창 안내문(placeholder) ════════
    PLACEHOLDER = "대화를 입력해 주세요."

    def _set_placeholder(self):
        self.inp.delete("1.0", "end")
        self.inp.insert("1.0", self.PLACEHOLDER)
        self.inp.configure(fg=C_DIM)
        self._ph_active = True

    def _clear_placeholder(self, e=None):
        if getattr(self, "_ph_active", False):
            self.inp.delete("1.0", "end")
            self.inp.configure(fg=C_TEXT)
            self._ph_active = False

    def _restore_placeholder(self, e=None):
        if not self.inp.get("1.0", "end").strip():
            self._set_placeholder()

    # ════════ 전송/스트리밍 ════════
    def on_return(self, event):
        self.send()
        return "break"

    def send(self):
        if getattr(self, "_ph_active", False):
            return
        text = self.inp.get("1.0", "end").strip()
        if not text or self.streaming:
            return
        self.inp.delete("1.0", "end")
        self._dispatch(text)

    def _dispatch(self, content, display=None):
        self._append("\n나 ▸ ", "user")
        self._append((display or content) + "\n", "bot")
        self.history.append({"role": "user", "content": content})
        self._cancel = False
        self.streaming = True
        self._set_stop_button()
        self._append("\nAI ▸\n", "label")
        self.chat.see("end")          # 새 턴은 항상 맨 아래로 (이후 스트리밍은 스마트 추적)
        self._ans_start = self.chat.index("end-1c")  # 답변 시작 위치 기록
        self._first_chunk = True
        self._load_i = 0
        self._answer_acc = []
        self._retry_msg = None
        threading.Thread(target=self._worker, daemon=True).start()
        self.after(40, self._poll)
        self._animate_loading()   # 점멸 로딩 표시 시작

    def _animate_loading(self):
        """첫 글자가 오기 전까지 '입력중 …'을 점멸 애니메이션으로 표시."""
        if not getattr(self, "_first_chunk", False) or not self.streaming:
            return
        base = getattr(self, "_retry_msg", None) or "입력중"
        frames = [base, base + " ·", base + " · ·", base + " · · ·"]
        follow = self._at_bottom()
        self.chat.configure(state="normal")
        self.chat.delete(self._ans_start, "end-1c")
        self.chat.insert("end", frames[self._load_i % len(frames)], "sys")
        if follow:
            self.chat.see("end")
        self.chat.configure(state="disabled")
        self._load_i += 1
        self.after(400, self._animate_loading)

    def _worker(self):
        for piece in core.stream_answer(self.client, self.model_key, self.mode, self.history,
                                        should_cancel=lambda: self._cancel):
            if isinstance(piece, tuple) and piece[0] == "__error__":
                self.q.put(("err", piece[1]))
            elif isinstance(piece, tuple) and piece[0] == "__retry__":
                self.q.put(("retry", piece[1:]))   # (대기초, 회차, 총회차)
            else:
                self.q.put(("chunk", piece))
        self.q.put(("done", None))

    def _poll(self):
        try:
            while True:
                kind, data = self.q.get_nowait()
                if kind == "chunk":
                    if self._first_chunk:        # 첫 글자 도착 → 대기표시 제거
                        self._retry_msg = None   # 재시도 안내 해제
                        self.chat.configure(state="normal")
                        self.chat.delete(self._ans_start, "end-1c")
                        self.chat.configure(state="disabled")
                        self._first_chunk = False
                    self._answer_acc.append(data)        # 원본 보관
                    self._append(data.replace("**", ""), "bot")   # 스트리밍 중 ** 제거
                elif kind == "retry":
                    wait, attempt, total = data
                    self._retry_msg = (
                        f"요청이 많아 약 {wait}초 후 자동으로 다시 시도합니다 (재시도 {attempt}/{total})"
                    )
                elif kind == "err":
                    if self._first_chunk:        # 대기표시 제거
                        self.chat.configure(state="normal")
                        self.chat.delete(self._ans_start, "end-1c")
                        self.chat.configure(state="disabled")
                        self._first_chunk = False
                    self._append(f"\n⚠️ {core.friendly_error(data)}\n", "err")
                elif kind == "done":
                    self._finish()
                    return
        except queue.Empty:
            pass
        if self.streaming:
            self.after(40, self._poll)

    def _finish(self):
        self.streaming = False
        self._set_send_button()
        answer = "".join(self._answer_acc)
        if not answer:
            if self._cancel:                      # 첫 글자 전에 중지
                self._sys("답변을 중지했습니다.")
                self._cancel = False
            return
        self.history.append({"role": "assistant", "content": answer})
        # 스트리밍된 원본(마크다운 기호 포함)을 지우고 깔끔하게 다시 렌더
        self.chat.configure(state="normal")
        self.chat.delete(self._ans_start, "end-1c")
        self.chat.configure(state="disabled")
        self._render_markdown(answer)
        if self._cancel:                          # 사용자가 중간에 멈춤 → 부분 답변 유지
            self._append("\n· 답변을 중지했습니다.\n", "warn")
            self._cancel = False
            return
        if self.mode in ("design", "review"):
            self._append(prompts.CITATION_NOTE + "\n", "warn")
        cjk = core.check_cjk(answer)
        if cjk:
            self._append("\n⚠️ 한자/가나 의심 문자 발견 (한글 교체 권장):\n", "warn")
            for item in cjk[:10]:
                self._append("   " + item + "\n", "warn")
            if len(cjk) > 10:
                self._append(f"   …외 {len(cjk) - 10}개\n", "warn")
        if self.mode in ("design", "review"):
            ts = datetime.now().strftime("%Y-%m-%d_%H%M")
            label = prompts.MODE_LABELS[self.mode]
            path = os.path.join(OUT_DIR, f"{ts}_{label}.md")
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(f"# [{label}] {ts}\n\n{answer}\n")
                self._sys(f"📄 저장됨: {path}")
            except Exception:
                pass
        self._append("\n✅ 답변이 완료되었습니다.\n", "sys")
        self._notify("답변 완료", "연구설계 챗봇 답변이 완료되었습니다.")

    def _notify(self, title, message):
        """답변/작업 완료를 알린다 — 소리(벨) + macOS 알림 센터."""
        try:
            self.bell()                       # 모든 OS 공통 소리
        except Exception:
            pass
        try:
            if platform.system() == "Darwin":
                safe_t = title.replace('"', "'")
                safe_m = message.replace('"', "'")
                subprocess.run(
                    ["osascript", "-e",
                     f'display notification "{safe_m}" with title "{safe_t}" sound name "Glass"'],
                    check=False, capture_output=True,
                )
        except Exception:
            pass

    # ════════ 출력 헬퍼 ════════
    def _at_bottom(self):
        """대화창 뷰가 맨 아래(근처)인지 — True면 새 내용을 따라 내려간다."""
        try:
            return self.chat.yview()[1] >= 0.999
        except Exception:
            return True

    def _on_chat_scroll(self, lo, hi):
        """스크롤바 위치 갱신."""
        self._sb.set(lo, hi)

    def _enable_copy(self, widget):
        """state=disabled인 Text에서도 드래그 선택·복사·전체 선택을 허용한다.
        편집(붙여넣기·삭제)은 막아 읽기 전용을 유지한다."""
        # 복사 단축키 (⌘C / Ctrl+C) 및 전체 선택 (⌘A / Ctrl+A)
        widget.bind("<Command-c>", lambda e: self._copy_selection(widget))
        widget.bind("<Control-c>", lambda e: self._copy_selection(widget))
        widget.bind("<Command-a>", lambda e: self._select_all(widget))
        widget.bind("<Control-a>", lambda e: self._select_all(widget))

        # 우클릭(맥은 Button-2, 그 외 Button-3) 컨텍스트 메뉴
        menu = tk.Menu(widget, tearoff=0)
        menu.add_command(label="복사", command=lambda: self._copy_selection(widget))
        menu.add_command(label="전체 선택", command=lambda: self._select_all(widget))

        def popup(e):
            try:
                menu.tk_popup(e.x_root, e.y_root)
            finally:
                menu.grab_release()
            return "break"

        widget.bind("<Button-2>", popup)
        widget.bind("<Button-3>", popup)

    def _copy_selection(self, widget):
        try:
            sel = widget.get("sel.first", "sel.last")
        except tk.TclError:
            sel = ""
        if sel:
            self.clipboard_clear()
            self.clipboard_append(sel)
        return "break"

    def _select_all(self, widget):
        widget.tag_add("sel", "1.0", "end-1c")
        return "break"

    def _append(self, text, tag):
        follow = self._at_bottom()              # 삽입 전 위치 판단
        self.chat.configure(state="normal")
        self.chat.insert("end", text, tag)
        if follow:                              # 바닥에 있었을 때만 따라감
            self.chat.see("end")
        self.chat.configure(state="disabled")

    def _render_markdown(self, text):
        """마크다운 기호를 실제 서식으로 바꿔 깔끔하게 출력 (** → 굵게, # → 제목, - → •)."""
        follow = self._at_bottom()
        self.chat.configure(state="normal")
        for raw in text.split("\n"):
            line = raw.rstrip()
            m = re.match(r"^(#{1,6})\s*(.*)", line)
            if m:                                   # 제목
                self._inline(m.group(2), tag="h"); self.chat.insert("end", "\n"); continue
            if re.match(r"^\s*([-*_])(\s*\1){2,}\s*$", line):  # 구분선
                self.chat.insert("end", "─" * 30 + "\n", "hr"); continue
            mb = re.match(r"^(\s*)[-*]\s+(.*)", line)
            if mb:                                  # 글머리
                self.chat.insert("end", mb.group(1) + "• ", "bot")
                self._inline(mb.group(2)); self.chat.insert("end", "\n"); continue
            self._inline(line); self.chat.insert("end", "\n")
        if follow:
            self.chat.see("end")
        self.chat.configure(state="disabled")

    def _inline(self, text, tag="bot"):
        """줄 안의 **굵게**를 처리하고 남은 마크다운 기호(** ` )는 제거."""
        text = text.replace("`", "")
        for part in re.split(r"(\*\*.+?\*\*)", text):
            if len(part) >= 4 and part.startswith("**") and part.endswith("**"):
                self.chat.insert("end", part[2:-2], "b" if tag == "bot" else tag)
            else:
                # 짝이 안 맞아 남은 ** 기호 제거
                self.chat.insert("end", part.replace("**", ""), tag)

    def _sys(self, text):
        self._append("\n· " + text + "\n", "sys")

    def _warn(self, text):
        self._append("\n⚠️ " + text + "\n", "warn")

    # ════════ 전역 단축키 (기본 꺼짐 — macOS 크래시 방지) ════════
    # 켜려면 환경변수 RC_HOTKEY=1 로 실행 + macOS 입력 모니터링 권한 허용.
    def setup_hotkey(self):
        if os.environ.get("RC_HOTKEY") != "1":
            return
        try:
            from pynput import keyboard
        except Exception:
            return

        def toggle():
            self.after(0, self._toggle_window)
        try:
            combo = "<cmd>+<shift>+r" if platform.system() == "Darwin" else "<ctrl>+<shift>+r"
            self._hk = keyboard.GlobalHotKeys({combo: toggle})
            self._hk.daemon = True
            self._hk.start()
        except Exception:
            pass

    def _toggle_window(self):
        if not self.winfo_viewable():
            self.deiconify(); self.lift(); self.focus_force()
        else:
            self.withdraw()


if __name__ == "__main__":
    App().mainloop()
